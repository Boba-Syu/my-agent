"""
Word 处理器

处理Word文档(.docx)的解析和分块。
"""

from __future__ import annotations

import logging
from typing import Any

from app.domain.rag.document_chunk import DocumentChunk
from app.domain.rag.document_processor import DocumentProcessor

logger = logging.getLogger(__name__)


class WordProcessor(DocumentProcessor):
    """
    Word 处理器
    
    使用 python-docx 解析Word文档。
    """
    
    @property
    def supported_types(self) -> list[str]:
        """支持的文件类型"""
        return ["docx"]
    
    def can_process(self, file_path: str) -> bool:
        """检查是否能处理"""
        return file_path.lower().endswith(".docx")
    
    def extract_text(self, file_path: str) -> str:
        """
        提取Word文本内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            提取的文本
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            text_parts = []
            
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("python-docx未安装，请运行: uv add python-docx")
            raise
        except Exception as e:
            logger.error(f"解析Word失败: {file_path}, 错误: {e}")
            raise ValueError(f"无法解析Word文档: {file_path}")
    
    def split_into_chunks(
        self,
        content: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ) -> list[DocumentChunk]:
        """
        将文本分块
        
        继承文本处理器的分块逻辑。
        """
        # 复用文本处理器的逻辑
        from app.infrastructure.rag.processors.text_processor import TextProcessor
        
        text_processor = TextProcessor()
        return text_processor.split_into_chunks(content, chunk_size, chunk_overlap)
    
    def process(
        self,
        file_path: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ) -> tuple[str, list[DocumentChunk]]:
        """处理Word文档"""
        text = self.extract_text(file_path)
        chunks = self.split_into_chunks(text, chunk_size, chunk_overlap)
        return text, chunks
